from pathlib import Path
import sys

sys.path.insert(0, str(Path("outputs/doc_vendor").resolve()))

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("outputs/SciData_Agent_taskbook_video_requirements_v2.docx")


def set_east_asia(run, font="Microsoft YaHei"):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def borders(table, color="D9E2EC"):
    tbl_pr = table._tbl.tblPr
    tbl_borders = tbl_pr.first_child_found_in("w:tblBorders")
    if tbl_borders is None:
        tbl_borders = OxmlElement("w:tblBorders")
        tbl_pr.append(tbl_borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        element = tbl_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tbl_borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def table_widths(table, widths):
    for row in table.rows:
        for i, width in enumerate(widths):
            cell = row.cells[i]
            cell.width = Inches(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")


def cell_text(cell, text, bold=False, size=8.8, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(str(text))
    r.bold = bold
    r.font.size = Pt(size)
    set_east_asia(r)
    if color:
        r.font.color.rgb = RGBColor.from_string(color)


def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.keep_with_next = True
    for r in p.runs:
        set_east_asia(r)
        r.font.color.rgb = RGBColor(46, 116, 181) if level == 1 else RGBColor(31, 77, 120)
    return p


def para(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(text)
    r.font.size = Pt(10.5)
    set_east_asia(r)
    return p


def bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.12
        r = p.add_run(item)
        r.font.size = Pt(10)
        set_east_asia(r)


def numbers(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.12
        r = p.add_run(item)
        r.font.size = Pt(10)
        set_east_asia(r)


def callout(doc, title, body, fill="F4F6F9", accent="1F4D78"):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    borders(table, "D7DEE8")
    table_widths(table, [6.5])
    cell = table.cell(0, 0)
    shade(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(10.5)
    r.font.color.rgb = RGBColor.from_string(accent)
    set_east_asia(r)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.12
    r2 = p2.add_run(body)
    r2.font.size = Pt(9.5)
    set_east_asia(r2)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def matrix(doc, headers, rows, widths, header_fill="E8EEF5"):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell_text(table.rows[0].cells[i], h, bold=True, size=9, color="0B2545")
        shade(table.rows[0].cells[i], header_fill)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cell_text(cells[i], value)
    borders(table)
    table_widths(table, widths)
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def code(doc, text):
    table = doc.add_table(rows=1, cols=1)
    borders(table, "DADCE0")
    table_widths(table, [6.5])
    cell = table.cell(0, 0)
    shade(cell, "F7F7F7")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    for i, line in enumerate(text.splitlines()):
        if i:
            p.add_run("\n")
        r = p.add_run(line)
        r.font.name = "Consolas"
        r.font.size = Pt(8.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    for style_name in ("Normal", "Heading 1", "Heading 2", "Heading 3", "List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    doc.styles["Normal"].font.size = Pt(10.5)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("SciData Agent 项目任务书")
    r.bold = True
    r.font.size = Pt(22)
    r.font.color.rgb = RGBColor(11, 37, 69)
    set_east_asia(r)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("基于赛道二讲解视频修订：多源科学数据查找、解析与整合 Data Agent")
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(85, 85, 85)
    set_east_asia(r)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("版本：V2.0 | 修订日期：2026-07-08 | 后续开发基准文档")
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(85, 85, 85)
    set_east_asia(r)

    callout(
        doc,
        "本文档用途",
        "本任务书用于替代旧版功能计划书，作为后续开发、测试、答辩材料和前端设计的统一依据。所有后续工作均围绕 Data Agent 闭环展开，而不是围绕单一 PDF 解析工具展开。",
        "EEF4FF",
    )

    heading(doc, "1. 视频要求提炼")
    para(doc, "赛题归属：挑战杯中国青年科技创新“揭榜挂帅”擂台赛，阿里云榜题，赛道二“数据场景”，AI Scientist 方向中的“科学知识与数据”。")
    para(doc, "视频中与本项目最相关的选题为“多源科学数据查找与整合”。该选题的本质不是构建普通解析器，而是构建 Data Agent：把论文、补充材料、数据库等来源中的科研信息转成可信、结构化、可追溯的数据。")
    bullets(doc, [
        "核心流程：来源发现 -> 文献/PDF 解析 -> 字段抽取 -> schema 对齐 -> 冲突校验 -> CSV/JSON/KG 输出。",
        "关键提醒：OCR 和 PDF parser 只是入口能力，不能替代数据智能体。",
        "关键目标：来源可信、字段准确、语义一致、可追溯、可复用、可校验。",
        "评审期待：系统必须能形成闭环，能够从原始科研资料到结构化数据，再到证据、日志和质量评价。",
    ])

    heading(doc, "2. 项目定位")
    para(doc, "项目名称：SciData Agent：基于通义千问的多源科学数据查找、解析与整合智能体。")
    para(doc, "一句话定位：本项目构建一个面向科研人员的 Data Agent，基于阿里云百炼通义千问大模型，将论文、补充材料、表格和数据库中的非结构化科研信息自动转化为带证据链的结构化数据，并支持字段归一、冲突校验、质量评估和 CSV/JSON 导出。")
    para(doc, "产品形态：后端 Agent 是核心，前端是交互入口。系统最终应包括可调用 API、交互式前端、代表性测试案例、处理日志、结构化输出、字段说明和技术报告。")

    heading(doc, "3. 建设原则")
    bullets(doc, [
        "必须是真实 Agent：必须由通义千问/百炼模型参与任务规划、语义抽取、质量判断或冲突解释，不能只做规则脚本。",
        "不能输出假数据：所有结构化记录必须来自输入文件、数据库或真实检索结果，并保留来源。",
        "证据优先：每个关键字段必须尽量绑定 source_file、page、evidence_text、confidence 和 warnings。",
        "结构化优先：输出自然语言总结可以有，但主结果必须是 CSV/JSON/KG 等可复用数据。",
        "闭环优先：系统要包含规划、抽取、规范化、校验、反馈和导出，而不是单点工具。",
        "最小可用优先：先实现稳定 MVP，再扩展 OCR、图表解析、外部数据库和知识图谱。",
    ])

    heading(doc, "4. 总体架构")
    para(doc, "系统采用“前端交互层 + API 服务层 + Agent 编排层 + 工具层 + 数据与证据层”的结构。")
    matrix(doc, ["层级", "职责", "最小实现", "后续增强"], [
        ["前端交互层", "输入问题、上传文件、查看结果和证据", "暂缓；先通过 CLI/API 使用", "React/Vue 页面、任务状态、证据面板、下载按钮"],
        ["API 服务层", "对外提供任务提交、查询和导出接口", "FastAPI 基础接口", "异步任务队列、用户空间、鉴权"],
        ["Agent 编排层", "规划任务、选择工具、组织抽取与校验", "Qwen planner + extractor + validator", "多轮反思、工具选择、失败重试、专家反馈闭环"],
        ["工具层", "解析 PDF/CSV/Excel，做规范化和导出", "pypdf、pandas、规则校验", "OCR、版面解析、表格识别、网页检索、数据库连接"],
        ["数据与证据层", "保存结构化结果、证据、日志和质量报告", "JSON/CSV/log 文件", "SQLite/PostgreSQL、向量库、知识图谱"],
    ], [1.1, 1.75, 1.75, 1.9])

    heading(doc, "5. Agent 工作流")
    numbers(doc, [
        "任务理解：接收研究问题和文件列表，识别科研领域、目标实体、指标字段和输出 schema。",
        "来源解析：解析 PDF、CSV、Excel 等输入，生成带来源定位的 text blocks 和 table blocks。",
        "抽取规划：由 Qwen 生成字段抽取计划，明确字段名称、别名、单位、必填/可选要求和质量规则。",
        "结构化抽取：由 Qwen 从文本块和表格块中抽取 records，不允许凭空补值。",
        "Schema 对齐：把不同来源的字段名称、同义词、单位和实体名称映射到统一 schema。",
        "质量校验：检查证据是否包含抽取值、字段是否缺失、单位是否合理、置信度是否足够。",
        "冲突检测：识别同一实体同一指标在不同来源中的不一致，并输出冲突说明和证据。",
        "结果导出：生成 result.json、result.csv、processing_log.json 和 quality_report.json。",
        "反馈迭代：根据 warnings、人工反馈或专家规则，修正 schema、抽取提示词和验证规则。",
    ])

    heading(doc, "6. 功能范围")
    matrix(doc, ["模块", "必须实现", "MVP 标准", "验收标准"], [
        ["文件输入", "PDF、CSV、Excel 上传/路径输入", "CLI/API 支持本地文件路径", "能读取至少 1 个 PDF 和 1 个 CSV/Excel"],
        ["任务规划", "Qwen 生成抽取计划和字段 schema", "记录 planner 输出", "日志中可见模型调用和字段计划"],
        ["文本解析", "按页解析 PDF 文本并保留页码", "支持 max_pdf_pages", "每个 text block 有 file/page/text"],
        ["表格解析", "解析 CSV/Excel 行列结构", "支持 pandas 读取", "表格字段可进入抽取流程"],
        ["字段抽取", "Qwen 输出结构化 records", "JSON schema 约束", "每条记录有 entity/field/value/source/evidence"],
        ["规范化", "字段名、单位、实体名归一", "内置常见指标规则", "同义字段合并到统一 field_name"],
        ["证据链", "绑定文件、页码、证据文本", "证据缺失时 warning", "关键数值必须能在 evidence_text 中找到或被标警告"],
        ["质量校验", "缺失、冲突、低置信度、单位异常检测", "Qwen validator + 规则校验", "输出 quality_report 或 warnings"],
        ["导出", "CSV/JSON/日志", "本地输出目录", "文件可被 Excel/程序直接读取"],
        ["API", "任务提交、查询、下载", "FastAPI 最小实现", "可用 curl/Postman 调用"],
        ["前端", "上传、任务状态、结果表、证据面板", "后置开发", "答辩前完成可交互页面"],
    ], [1.05, 1.95, 1.65, 1.85])

    heading(doc, "7. 数据结构与输出字段")
    para(doc, "MVP 采用通用科研记录结构，既能覆盖材料科学，也能覆盖机器学习论文。具体字段可由 Agent 根据研究问题动态扩展。")
    code(doc, '''{
  "task_id": "task_xxx",
  "question": "用户研究问题",
  "domain": "material science / machine learning / biology / other",
  "records": [
    {
      "entity": "材料、模型、实验对象或数据集",
      "field_name": "统一后的字段名",
      "field_value": "字段值",
      "unit": "单位；无量纲指标标记为 dimensionless",
      "condition": "实验条件或上下文",
      "source_file": "来源文件名",
      "page": "页码或表格行号",
      "evidence_text": "原文证据",
      "confidence": 0.0,
      "warnings": ["证据缺失、单位缺失、冲突等"]
    }
  ],
  "quality_report": {
    "record_count": 0,
    "warning_count": 0,
    "conflicts": []
  }
}''')

    heading(doc, "8. MVP 开发计划")
    matrix(doc, ["阶段", "目标", "具体任务", "完成标志"], [
        ["阶段 1：后端闭环", "把现有 Agent 改成符合视频要求的 Data Agent MVP", "完善 planner/extractor/validator；补 quality_report；强化证据检查；优化字段归一", "CLI 跑通 PDF + CSV 案例，输出四类文件"],
        ["阶段 2：质量提升", "减少幻觉和错误抽取", "数值必须出现在证据中；无量纲单位规则；去重；冲突检测；测试集", "错误值被过滤或标记，冲突能被报告"],
        ["阶段 3：API 完整化", "形成可调用接口", "POST /api/analyze；GET /api/tasks/{id}；GET /api/tasks/{id}/export", "Postman/curl 可提交任务并下载结果"],
        ["阶段 4：前端最小版", "满足交互式前端要求", "问题输入、文件上传、状态、结果表、证据侧栏、下载", "浏览器可完成一次端到端演示"],
        ["阶段 5：参赛材料", "准备答辩和提交材料", "PPT/PDF 20 页内；测试案例；技术报告；源码说明；演示视频", "材料齐全且与官方要求一致"],
    ], [1.25, 1.35, 2.55, 1.35])

    heading(doc, "9. 具体待办清单")
    para(doc, "以下待办是后续开发的主线，建议按优先级推进。")
    heading(doc, "9.1 P0：必须完成", 2)
    bullets(doc, [
        "重写/强化 Planner Prompt：明确“来源发现、字段抽取、schema 对齐、冲突校验、CSV/JSON/KG 输出”这条链路。",
        "增加 quality_report.json：记录总记录数、警告数、冲突数、字段覆盖率、证据覆盖率、模型调用情况。",
        "实现证据强校验：如果 metric/value 不在 evidence_text 中，必须 warning、降置信度，必要时从最终结果过滤。",
        "实现 dimensionless 单位规则：FID、KID、SSIM、LPIPS、accuracy、RMSE 等按指标类型规范 unit。",
        "实现基础冲突检测：同一 entity + field_name 出现多个不同 field_value 时生成 conflict 记录。",
        "完善 README 使用说明：环境、API key、CLI、API、输出文件解释、常见问题。",
        "准备 3 个代表性测试案例：PDF 单源、CSV/Excel 表格、PDF + 表格多源整合。",
    ])
    heading(doc, "9.2 P1：重要增强", 2)
    bullets(doc, [
        "增加 PDF 表格解析能力，优先尝试 pdfplumber；识别失败时保留日志。",
        "增加字段别名词典，例如 PCE/efficiency/光电转换效率、stability/稳定性、dataset/数据集。",
        "增加去重逻辑：同一来源同一证据重复抽取时合并。",
        "增加领域 schema 模板：材料科学、机器学习、生物医学三个模板。",
        "增加处理日志可视化字段：每个工具调用、输入块数量、模型返回记录数、校验结果。",
        "补 API 单元测试和端到端测试。",
    ])
    heading(doc, "9.3 P2：展示与竞争力增强", 2)
    bullets(doc, [
        "实现网页检索/来源发现模块，可接入学术搜索或用户提供 URL。",
        "实现 OCR 或图片图表解析，覆盖扫描 PDF 和论文图表。",
        "实现知识图谱导出 KG：至少输出 nodes.json 和 edges.json。",
        "实现人工反馈闭环：用户可确认/修改记录，系统记录反馈并重新校验。",
        "做前端证据面板：点击表格行时展示原文证据、页码和 warning。",
        "准备 10 分钟以内演示视频，展示端到端流程和结果可信性。",
    ])

    heading(doc, "10. 测试计划")
    matrix(doc, ["测试类型", "输入", "检查点", "通过标准"], [
        ["PDF 单源抽取", "一篇可复制文本 PDF", "页码、证据、字段值、日志", "抽取记录非空，关键值有证据"],
        ["CSV/Excel 表格抽取", "结构化数据表", "字段映射、单位、导出", "输出 CSV/JSON 与原表一致或合理归一"],
        ["PDF + 表格整合", "同一主题论文和表格", "去重、schema 对齐、多源证据", "同类字段合并为统一 schema"],
        ["冲突检测", "人为构造或真实矛盾数据", "conflicts 输出", "能列出冲突值和来源证据"],
        ["错误提示", "不相关文件或字段缺失", "warnings 和质量报告", "不会编造数据，能说明未找到"],
        ["API 测试", "curl/Postman 上传文件", "任务状态和导出", "接口可稳定返回结果文件"],
    ], [1.25, 1.65, 1.9, 1.7])

    heading(doc, "11. API 与前端规划")
    heading(doc, "11.1 API", 2)
    code(doc, '''POST /api/analyze
  multipart/form-data: question, files[], max_pdf_pages
  return: task_id, status

GET /api/tasks/{task_id}
  return: status, summary, records, quality_report, processing_log

GET /api/tasks/{task_id}/export?format=json|csv
  return: result file

GET /api/health
  return: qwen_configured, parser_status, version''')
    heading(doc, "11.2 前端", 2)
    bullets(doc, [
        "首页即工作台，不做营销落地页。",
        "左侧：研究问题输入、文件上传、参数设置、开始分析按钮。",
        "中间：任务状态、结果表格、字段筛选、冲突/警告标记。",
        "右侧：证据面板，展示来源文件、页码、原文证据和模型解释。",
        "顶部或底部：导出 CSV/JSON、查看日志、查看质量报告。",
    ])

    heading(doc, "12. 参赛提交材料映射")
    matrix(doc, ["官方/视频期待", "项目对应产物", "准备状态"], [
        ["使用阿里云模型能力", "百炼/通义千问 API 配置、模型调用日志", "已具备基础能力，需在报告中明确"],
        ["可调用 API", "FastAPI analyze/tasks/export 接口", "需完善和测试"],
        ["交互式前端", "上传、结果表、证据面板", "后续开发"],
        ["代表性测试案例", "PDF、CSV/Excel、多源整合、冲突检测", "需整理固定案例"],
        ["结构化输出示例", "result.json、result.csv、字段说明", "已具备雏形，需规范"],
        ["来源列表/处理日志", "processing_log.json、quality_report.json", "需强化 quality_report"],
        ["技术报告和源码", "README、架构图、核心模块说明", "需整理"],
        ["演示视频", "10 分钟以内端到端演示", "可选但建议准备"],
    ], [1.75, 2.95, 1.45])

    heading(doc, "13. 风险与边界")
    bullets(doc, [
        "PDF 解析质量风险：扫描件、复杂表格、双栏论文和图表会降低抽取质量，需要 OCR/版面解析增强。",
        "LLM 幻觉风险：必须通过证据强校验、置信度和 warnings 降低风险。",
        "领域泛化风险：MVP 先支持通用字段和少量领域模板，不承诺覆盖所有学科。",
        "外部检索风险：网页和数据库接口稳定性不可控，MVP 可先支持用户上传来源。",
        "评审表达风险：答辩时必须强调 Data Agent 闭环，避免被认为只是 PDF parser。",
    ])

    heading(doc, "14. 下一步执行顺序")
    numbers(doc, [
        "先改后端 Agent：补 quality_report、证据强校验、冲突检测和 schema 对齐。",
        "用真实 PDF 和构造 CSV 跑端到端测试，保存稳定输出作为展示案例。",
        "整理 API 文档和 curl 示例，保证前端开发前接口稳定。",
        "开发前端最小工作台，围绕结果表和证据链展示。",
        "把架构、流程、测试结果整理成 20 页以内 PPT/PDF。",
    ])
    callout(
        doc,
        "最终判断",
        "本项目应该被定义为“科研数据整合 Data Agent”，而不是“论文解析器”。后续所有功能、测试和答辩材料都要围绕“多源输入 -> Agent 规划 -> 结构化抽取 -> schema 对齐 -> 冲突校验 -> 证据链 -> CSV/JSON/KG 输出”的闭环展开。",
        "FFF7E6",
        "7A5A00",
    )

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run("SciData Agent 项目任务书 V2.0")
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor(120, 120, 120)
    set_east_asia(r)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT.resolve())


if __name__ == "__main__":
    build()
